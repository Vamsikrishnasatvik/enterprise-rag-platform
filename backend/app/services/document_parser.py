from pathlib import Path

import fitz
from docx import Document as DocxDocument

from app.schemas.parser import ParsedDocument


def parse_document(
    file_path: str,
    file_type: str,
) -> ParsedDocument:
    """
    Dispatch parser based on content type or extension.
    """

    file_type = (file_type or "").lower()
    suffix = Path(file_path).suffix.lower()

    if (
        file_type == "application/pdf"
        or suffix == ".pdf"
    ):
        return parse_pdf(file_path)

    if (
        file_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == ".docx"
    ):
        return parse_docx(file_path)

    if (
        file_type.startswith("text/")
        or suffix == ".txt"
        or suffix == ".csv"
    ):
        return parse_txt(file_path)

    raise ValueError(
        f"Unsupported document type: {file_type}"
    )


def parse_pdf(
    file_path: str,
) -> ParsedDocument:
    """
    Parse PDF using PyMuPDF.
    """

    document = fitz.open(file_path)

    pages = []
    full_text = []

    for page in document:
        text = page.get_text("text")
        pages.append(text)
        full_text.append(text)

    text = "\n".join(full_text)

    metadata = {
        "filename": Path(file_path).name,
        "content_type": "application/pdf",
        "pages": pages,
    }

    return ParsedDocument(
        text=text,
        page_count=len(pages),
        metadata=metadata,
    )


def parse_docx(
    file_path: str,
) -> ParsedDocument:
    """
    Parse DOCX.
    """

    document = DocxDocument(file_path)

    paragraphs = [
        p.text
        for p in document.paragraphs
        if p.text.strip()
    ]

    text = "\n".join(paragraphs)

    metadata = {
        "filename": Path(file_path).name,
        "content_type": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    }

    return ParsedDocument(
        text=text,
        page_count=1,
        metadata=metadata,
    )


def parse_txt(
    file_path: str,
) -> ParsedDocument:
    """
    Parse TXT.
    """

    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore",
    ) as f:
        text = f.read()

    metadata = {
        "filename": Path(file_path).name,
        "content_type": "text/plain",
    }

    return ParsedDocument(
        text=text,
        page_count=1,
        metadata=metadata,
    )